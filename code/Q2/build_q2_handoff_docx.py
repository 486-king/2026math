from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "Q2_修订版_编程手与论文手工作文档.docx"
FIGURE = (
    ROOT
    / "results"
    / "Q2"
    / "experiments"
    / "round2"
    / "figures"
    / "q2_capacity_frontier.png"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CAUTION = "FFF4CE"
RED = "9B1C1C"
MUTED = "666666"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_font(run, size=11, bold=None, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for j, text in enumerate(headers):
        cell = table.rows[0].cells[j]
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, 9.5, True, DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for j, text in enumerate(row):
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(text))
            set_font(r, 9.2)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_numbering_definition(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [
        int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))
    ]
    abs_id = max(existing_abs, default=0) + 1
    num_id = max(existing_num, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abs_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_id = OxmlElement("w:abstractNumId")
    abstract_id.set(qn("w:val"), str(abs_id))
    num.append(abstract_id)
    numbering.append(num)
    return num_id


def numbered(doc, text, num_id):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)
    r = p.add_run(text)
    set_font(r)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def callout(doc, label, text, fill=LIGHT_GRAY, color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}　")
    set_font(r, 10.5, True, color)
    r = p.add_run(text)
    set_font(r, 10.5)
    set_table_geometry(table, [TABLE_WIDTH])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, 9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    r2 = paragraph.add_run(" 页")
    set_font(r2, 9, color=MUTED)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "2026 数学建模校赛｜B题 第2问执行手册"
    set_font(header.runs[0], 9, color=MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    programmer_num_id = add_numbering_definition(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("第2问：编程手与论文手任务清单")
    set_font(r, 24, True, DARK_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("面向第一次阅读题目的队员｜方案 Q2-O3 / Q2-A｜B 为强制 baseline")
    set_font(r, 12, color=MUTED)

    callout(
        doc,
        "先看结论",
        "在相对 M1/S1 模型中，1 枚弹最多完整遮蔽约 10.376 s，小于最坏探测窗口 "
        "25.361 s；一个两弹方案已被连续时间区间证书与独立几何方法共同验证，因此在允许锁定前预部署的相对模型中，完成全窗口防御的最少弹药数为 2。"
    )
    callout(
        doc,
        "必须先统一时间",
        "t=0 表示导弹进入 8000 m 边界并完成锁定，防御考核窗口为 [0,25.361] s。"
        "两弹全窗口方案的第一枚弹需要在 t=0 前接收指令、释放并起爆，所以结论依赖锁定前预警部署。"
        "如果规定所有无人机只能从 t=0 开始接受任务，则至少 [0,5.5) s 没有烟幕，不能宣称全窗口防御。",
        fill=CAUTION,
        color=RED,
    )

    heading(doc, "一、这道题到底要做什么", 1)
    doc.add_paragraph(
        "把舰船看成半径 80 m 的圆盘，把每一团烟幕看成随时间变化的圆盘。第二问不是只判断“舰船中心是否在烟里”，而是要求舰船整个圆盘在一段连续时间内都落在多团烟幕圆盘的并集里。"
    )
    bullet(doc, "优先目标：先判断能否覆盖完整的 M1 探测窗口；能覆盖时，再减少干扰弹数量。")
    bullet(doc, "能力目标：分别求 1、2、3 枚弹能够形成的最大连续完整遮蔽时长。")
    bullet(doc, "比较目标：用方案 B 的单弹区间接续给出透明、保守、可解释的下界。")
    bullet(doc, "禁止误判：不能只检查时间网格，也不能只检查舰船圆周或中心。")

    heading(doc, "二、已经确认的建模口径", 1)
    rows = [
        ("舰船", "半径 80 m 的圆盘；圆心沿 x 轴以 7.71 m/s 匀速运动。"),
        ("烟幕", "最大半径 120 m；前 18 s 半径不变，随后 5 s 线性衰减至 0。"),
        ("干扰弹 S1", "投放后 3.5 s 继承无人机投放瞬间水平速度；起爆后烟幕中心固定。"),
        ("导弹 M1", "速度大小恒定，方向实时指向舰船中心；最坏探测窗口 25.3610426206 s。"),
        ("风", "题面未给风向、风速，本轮不引入；漂移只作为后续鲁棒性参数。"),
        ("任务时钟", "t=0 为 8000 m 锁定；t_cmd→t_d=t_cmd+2→t_b=t_d+3.5。负时刻事件表示锁定前预任务。"),
        ("坐标", "当前结果是相对坐标；缺少无人机初始点和基地基准，绝对飞行可达性仍待输入。"),
    ]
    add_table(doc, ["对象", "统一口径"], rows, [1800, 7560])

    heading(doc, "三、当前可直接使用的数值结论", 1)
    rows = [
        ("1", "10.376134890", "10.376134890", "解析精确；不足以覆盖 M1"),
        ("2", "29.071871042", "20.752269780", "A 为共线前沿；贴近桥接边界"),
        ("3", "42.523129869", "31.128404669", "A 为当前最佳已验证共线解；非全局最优证明"),
    ]
    add_table(
        doc,
        ["弹数", "A：烟幕并集/s", "B：保守接续/s", "结论强度"],
        rows,
        [900, 2100, 2100, 4260],
    )
    callout(
        doc,
        "读表提示",
        "虚线门槛相当于 M1 最坏窗口 25.361 s：1 枚弹两种方法都不足；"
        "2 枚弹只有烟幕并集方案 A 超过门槛；3 枚弹的 A/B 都超过门槛，"
        "但 A 的空间互补能力明显更强。",
        fill=LIGHT_BLUE,
    )

    heading(doc, "四、应优先报告的两弹全窗口方案", 1)
    callout(
        doc,
        "统一到锁定时钟后的相对方案",
        "防御窗口为 [0,25.3610426206] s；烟幕中心 x=(33.73703319, 161.83285766) m；"
        "指令时刻 t_cmd=(-7.88582566, 2.94018905) s；实际投放时刻 "
        "t_d=(-5.88582566, 4.94018905) s；起爆时刻 "
        "t_b=(-2.38582566, 8.44018905) s。第一枚弹明确属于锁定前预任务。",
    )
    bullet(doc, "连续时间区间证书：覆盖 25.3610426206 s，9 个认证时间盒，0 个未决盒，0 个时间空档。")
    bullet(doc, "独立截面几何：最小平方截面裕度 1463.887280249688 m²，舰船边界和内部均无漏点。")
    bullet(doc, "高精度圆盘差集：关键时刻最大未覆盖面积为 0。")
    bullet(doc, "两种几何验证一致；这才是允许宣称“两弹全窗口可行”的证据。")
    bullet(doc, "绝对首投点和 12 km 半径检查仍阻塞，不得自行补造无人机初始位置。")

    heading(doc, "五、编程手任务清单（按顺序执行）", 1)
    callout(
        doc,
        "给编程手的话",
        "你不需要重新猜模型。先复现已有结果，再做误差和风漂扩展。每一步都要保存输入、输出和判定标准；不要只截图。",
        fill="EAF3E8",
    )
    steps = [
        "建立一个只读参数文件：录入舰速 7.71 m/s、舰半径 80 m、烟幕最大半径 120 m、18 s 恒定期、5 s 衰减期、3.5 s 起爆延迟、无人机速度 28 m/s。禁止修改原题 Word。",
        "统一任务时钟和相对坐标：t=0 为导弹在 8000 m 完成锁定，舰船圆心为 s(t)=7.71t；烟幕中心记为 c_j；所有长度统一用 m，时间统一用 s。负时刻命令、投放或起爆必须标注为锁定前预任务。",
        "实现烟幕半径函数：起爆前为 0；起爆后 0–18 s 为 120；18–23 s 线性降到 0；23 s 后为 0。为分段点写单元测试。",
        "实现“完整圆盘并集覆盖”判据。对 ξ∈[-80,80]，计算 L_j(ξ,t)=r_j²-80²-(s-c_j)²-2(s-c_j)ξ；固定 t 时 L_j 对 ξ 是直线，只需检查端点与两两交点处的 max_j L_j 最小值。",
        "先复现两弹全窗口方案。必须同时运行区间时间证书和独立截面验证；若任一方法失败或二者结论不一致，立即停止，不得输出“可行”。",
        "复现 baseline：单弹完整遮蔽上限 T1=2(120-80)/7.71；依次计算 1T1、2T1、3T1，并用时间区间并集检查接续是否有空档。",
        "复现能力前沿：1 弹用解析式；2 弹求桥接切触条件和末端衰减根；3 弹用固定随机种子的多起点搜索，再用精确截面与更细独立时间扫描淘汰隐藏空档。",
        "检查无人机相对转场：由投放点到烟幕中心的惯性位移为 28×3.5=98 m；逐对检查投放时间差至少 1 s，并检查 28Δt 是否不小于相邻投放点距离。再单独输出 absolute_execution_status：缺少 u_i(a_i) 和基地参考点时必须为 blocked。",
        "做鲁棒性实验：优先扰动舰速、烟幕半径、起爆延迟和烟幕中心漂移。两弹全窗口方案与容量前沿方案要分开测；前沿贴边失效不等于全窗口方案失效。",
        "整理交付物：参数 JSON、每个方案的时刻表 CSV、验证指标 JSON、容量前沿 CSV、诊断图、运行摘要和环境版本。所有脚本必须一条命令可重复运行。",
    ]
    for text in steps:
        numbered(doc, text, programmer_num_id)

    heading(doc, "六、编程手验收标准", 2)
    checks = [
        "同一版本连续运行两次，核心 JSON 完全一致。",
        "两弹全窗口方案：区间证书 0 个未决盒，独立验证无边界或内部漏点，两种方法结论一致。",
        "故意把两团烟幕拉开到明显断裂时，验证器必须返回 FAIL，不能把错误方案判成可行。",
        "1、2、3 弹容量单位均为秒，且明确区分“解析精确”“共线优化前沿”“当前最佳已验证”。",
        "三弹不得写成“全局最优”，除非另有可核查的匹配上界证明。",
        "缺失的无人机绝对初始状态必须输出 blocked，不得填 0 或假设基地在原点。",
        "若禁用锁定前预任务，程序必须返回初始 5.5 s 裸露证书，不能继续沿用两弹全窗口结论。",
    ]
    for text in checks:
        bullet(doc, "□ " + text)

    heading(doc, "七、论文手任务清单（按论文顺序写）", 1)
    callout(
        doc,
        "给论文手的话",
        "你的任务不是复制代码，而是把“为什么这样建、为什么验证可信、哪些结论有多强”写清楚。所有数值只能从已保存的结果文件取。",
        fill="EAF3E8",
    )
    paper_steps = [
        "问题重述：说明第二问同时包含最少弹药的全窗口防御和 1、2、3 枚弹的连续遮蔽能力比较；先满足 100% 防御，再减少资源。",
        "模型假设：逐条写舰船圆盘、烟幕圆盘、S1 惯性运动、起爆后中心固定、无风漂；把“题目给定”“必要假设”“简化假设”分开。",
        "符号说明：至少列 s(t)、R_s、c_j、r_j(t)、t_j^b、t_j^d、ξ、L_j、G_∪(t)，并给出单位。",
        "几何模型：先解释错误做法为何不够——舰船中心被遮住不代表整舰被遮住；单弹区间相加也会漏掉多烟幕空间互补。",
        "核心判据推导：从圆盘竖直截面平方高度出发，推到 L_j(ξ,t)；说明完整覆盖等价于 min_ξ max_j L_j(ξ,t)≥0。",
        "优化目标：用词典序表达，第一层最小化未覆盖时间，达到 0 后第二层最小化弹数；能力前沿再最大化单个连续完整遮蔽区间长度。",
        "两弹结果：先证明 1 枚弹的解析上限 10.376 s 小于 25.361 s，再引用两弹连续证书，从而得到“允许锁定前预部署的相对模型下最少 2 枚”的逻辑闭环。",
        "多弹增益：报告 A/B 对照表；强调 A 的提升来自烟幕并集的空间互补，不是把同类优化器换名字。",
        "验证章节：写明区间时间证书、精确截面法、圆盘差集、反例测试和重复运行。不能用“取了很多网格点”代替连续验证。",
        "局限性：三弹是当前最佳已验证共线方案，尚无全局上界；容量前沿贴边、对误差敏感；绝对首投点和 12 km 检查缺数据；风漂待扩展。",
        "任务时钟说明：明确写 t=0 为 8000 m 锁定，第一枚弹的负时刻事件是预任务；不得把负时刻藏在“整体平移”中。",
        "结论措辞：作战建议采用有正裕度的两弹全窗口方案；29.072 s 和 42.523 s 仅用于展示容量与协同增益。",
    ]
    paper_num_id = add_numbering_definition(doc)
    for text in paper_steps:
        numbered(doc, text, paper_num_id)

    heading(doc, "八、论文中建议使用的表和图", 2)
    rows = [
        ("表 1", "参数与符号表", "单位、来源、是否题面给定"),
        ("表 2", "两弹全窗口时刻表", "烟幕中心、投放/起爆时刻、覆盖窗口"),
        ("表 3", "A/B 能力前沿", "1、2、3 弹持续时间与结论强度"),
        ("表 4", "验证证据表", "区间盒、最小裕度、漏点面积、复跑一致性"),
        ("图 1", "能力前沿图", "A、B 与 M1 最坏窗口的横向比较"),
        ("图 2", "两弹几何示意", "关键时刻舰船圆盘与烟幕并集"),
    ]
    add_table(doc, ["编号", "内容", "必须表达的信息"], rows, [1000, 2700, 5660])

    heading(doc, "九、禁止出现的说法", 1)
    warnings = [
        "“两弹在 1000 个时间点都通过，所以连续时间一定通过。”",
        "“舰船中心位于烟幕内，所以整艘舰船被遮蔽。”",
        "“三弹 42.523 s 是全局最优。”（当前没有匹配上界证明）",
        "“无人机绝对首投点已确定。”（题目缺少初始位置和基地基准）",
        "“无人机都从 t=0 才开始行动，但两弹仍能从 t=0 全程遮蔽。”（与 5.5 s 事件延迟矛盾）",
        "“模型考虑了风。”（本轮没有风数据，也没有引入风漂）",
        "“保证国一”或“完美解决”。",
    ]
    for text in warnings:
        bullet(doc, "× " + text)

    heading(doc, "十、文件交接地图", 1)
    rows = [
        ("正式运行摘要", "results/Q2/experiments/round2/run_summary.json"),
        ("两弹最少资源方案", "results/Q2/experiments/round2/metrics/q2_two_bomb_minimum_resource_plan.json"),
        ("能力前沿", "results/Q2/experiments/round2/metrics/q2_capacity_frontier.json"),
        ("方案时刻表", "results/Q2/experiments/round2/tables/q2_verified_schedules.csv"),
        ("A/B 对照表", "results/Q2/experiments/round2/tables/q2_capacity_frontier.csv"),
        ("连续证书", "results/Q2/experiments/round1/metrics/q2_continuous_validation.json"),
        ("主运行脚本", "code/Q2/q2_run_round2.py"),
        ("代码审查", "code/Q2/reviews/q2_python_review_round2.json"),
    ]
    add_table(doc, ["用途", "相对路径"], rows, [2600, 6760])

    heading(doc, "十一、下一步最小行动", 1)
    callout(
        doc,
        "建议顺序",
        "编程手先复跑 round2 并补做参数扰动/风漂鲁棒性；论文手同时搭好第2问模型与验证章节骨架，但暂不把绝对首投点、风漂结果或三弹全局最优写成已知事实。",
        fill=CAUTION,
        color=RED,
    )

    doc.core_properties.title = "Q2 编程手与论文手任务清单"
    doc.core_properties.subject = "数学建模 B 题第2问执行与交接"
    doc.core_properties.author = "建模组"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
