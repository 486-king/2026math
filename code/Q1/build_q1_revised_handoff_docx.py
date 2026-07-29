from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "code" / "Q2"))
from build_q2_handoff_docx import (  # noqa: E402
    BLUE,
    CAUTION,
    DARK_BLUE,
    LIGHT_BLUE,
    LIGHT_GRAY,
    MUTED,
    RED,
    add_numbering_definition,
    add_page_number,
    add_table,
    bullet,
    callout,
    heading,
    numbered,
    set_font,
)


OUT = ROOT / "Q1_修订版_编程手与论文手分工.docx"


def add_formula(doc: Document, formula: str, note: str = "") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    from build_q2_handoff_docx import shade, set_table_geometry

    shade(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula)
    set_font(run, 10.2, False, DARK_BLUE, "Cambria Math")
    if note:
        run = p.add_run("　" + note)
        set_font(run, 9.2, False, MUTED)
    set_table_geometry(table, [9360])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "2026 数学建模校赛｜B题 Q1 修订执行手册"
    set_font(header.runs[0], 9, False, MUTED)
    add_page_number(section.footer.paragraphs[0])


def build() -> Path:
    doc = Document()
    configure_document(doc)
    num_id = add_numbering_definition(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("问题一修订版：编程手与论文手分工")
    set_font(run, 24, True, DARK_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(
        "依据队友《总体审查结论》修订｜保留原证明｜升级整题接口｜面向首次读题成员"
    )
    set_font(run, 11.5, False, MUTED)

    callout(
        doc,
        "修订后的总判断",
        "原来的数学核心没有错：在 G1 纯追踪、8000 m 处已经锁定、"
        "S1 固定云心、O0 二维完整圆盘覆盖、U0 无名义风漂条件下，"
        "单弹全探测窗口遮蔽严格不可行。需要改的是事件符号、输入状态、"
        "可达性层和 Q1→Q2 接口，而不是重新盲搜一个“唯一最优坐标”。",
    )

    heading(doc, "一、编程手和论文手先共同读懂这四句话", 1)
    bullet(
        doc,
        "第一问不是“程序没搜到解”，而是结构上界 10.3761 s 小于探测窗口下界 24.1677 s，因此全窗口可行集为空。",
    )
    bullet(
        doc,
        "不可行不等于没有成果：还要给出最大连续遮蔽补偿族，以及前段、中段、后段三类代表策略。",
    )
    bullet(
        doc,
        "10.3761 s 是结构理论上界，不自动等于无人机真正可执行的最优值；真实可执行值还要受初态、命令时刻、释放点和 12 km 约束筛选。",
    )
    bullet(
        doc,
        "Q2 不是把单弹时长乘以弹数，而是检验舰船圆盘是否被多团烟幕圆盘的并集完整覆盖。",
    )

    heading(doc, "二、队友建议的采纳结论", 1)
    rows = [
        ("直接采纳", "8000 m 处已锁定前提；G1/G2 重命名；三类时刻；98 m 标为模型假设；结构/执行分层；四类状态；Q1→Q2 联合覆盖接口"),
        ("条件采纳", "12 km 暂按相对初始起飞点的距离；无参考点时保持 blocked；风漂只给参数曲线"),
        ("保留原结果", "10.3761 s、24.1677 s、13.7916 s；A/B/C 三重证据；参数化补偿族"),
        ("不放入Q1主模型", "二阶覆盖与N−1韧性留给Q3；视线投影O1只作扩展；Lipschitz网格不替代解析/区间证书"),
    ]
    add_table(doc, ["处理", "具体内容"], rows, [1600, 7760])

    heading(doc, "三、所有人统一使用的新符号", 1)
    rows = [
        ("G1", "纯追踪导弹", "速度大小恒定，方向实时指向舰船中心"),
        ("G2", "固定航向对照", "只有窗口时长条件通过仍不能认定整体可行"),
        ("t_cmd", "投弹命令时刻", "主解释下 2 s 后实际释放"),
        ("t_d", "实际释放时刻", "t_d=t_cmd+2"),
        ("t_b", "起爆时刻", "t_b=t_d+3.5=t_cmd+5.5"),
        ("t_m", "覆盖区间中点", "替代原来会与 command 混淆的 t_c"),
        ("Delta(t)", "统一覆盖缺陷", "Delta≤0 才表示完整舰船圆盘被烟幕并集覆盖"),
        ("T_structural_max", "结构上界", "10.376134889753567 s"),
        ("T_executable_star", "可执行最优", "当前缺绝对场景输入，尚不能数值确定"),
    ]
    add_table(doc, ["符号", "含义", "必须记住"], rows, [1500, 2600, 5260])

    heading(doc, "四、三条不能改写的数学结论", 1)
    add_formula(
        doc,
        "T_structural_max = 2(R_c-R_s)/V_s = 10.3761348898 s",
        "单固定烟幕完整覆盖的结构上界",
    )
    add_formula(
        doc,
        "T_detect_lower = (8000-80)/(320+7.71) = 24.1677092551 s",
        "G1 且 8000 m 已锁定",
    )
    add_formula(
        doc,
        "T_naked_lower >= 24.1677092551-10.3761348898 = 13.7915743654 s",
        "单弹至少存在的总裸露时间下界",
    )
    callout(
        doc,
        "允许的正式结论",
        "在 G1+S1+O0+U0 及 8000 m 已锁定的标准场景中，"
        "单枚干扰弹无法在整个探测窗口持续完整遮蔽舰船。不要缩写成“问题一无解”。",
        fill="EAF3E8",
    )

    heading(doc, "五、编程手分工：你最终要交什么", 1)
    callout(
        doc,
        "编程手的核心职责",
        "维护可复现证据，不重新猜模型。先跑通现有四轮结果，再补场景化可达性层。"
        "任何真实初态缺失时都返回 blocked，不得自动填 0。",
        fill="EAF3E8",
    )

    heading(doc, "5.1 第一天先做的准备", 2)
    preparation = [
        "阅读原题第 1 问、results/Q1/reports/q1_teammate_review_revision.md 和 interfaces/Q1_to_Q2_coverage_contract.md；不要只看旧任务清单。",
        "确认 Python 环境能运行 NumPy、SciPy、pandas；记录版本，不在代码中写个人绝对路径。",
        "运行 q1_run.py、q1_interval_certificate.py、q1_parametric_compensation.py、q1_robustness.py、q1_architecture_upgrade.py。",
        "核对三项核心数字与本手册一致；若不一致，先停止并定位，不要继续画图或调优化器。",
    ]
    for item in preparation:
        numbered(doc, item, num_id)

    heading(doc, "5.2 程序模块逐项任务", 2)
    programmer_tasks = [
        (
            "P1 公共参数和烟幕半径",
            "维护 q1_common.py；所有单位统一为 m、s、m/s；检查烟幕龄期 0、18、23 s 的半径。",
            "端点分别得到 120、120、0 m；原题 Word 未被修改。",
        ),
        (
            "P2 命令—释放—起爆事件链",
            "用 t_cmd、t_d、t_b 三个字段；验证 2、3.5、5.5 s 三个差值；旧 drop_time 只能作为 release 的兼容别名。",
            "事件链测试 PASS；任何输出中不再用 t_c 表示命令或覆盖中点。",
        ),
        (
            "P3 统一覆盖缺陷",
            "保留 single_smoke_margin，并验证其等于 -Delta；多烟幕不得用有限网格冒充连续证明，必须调用 Q2 几何核。",
            "1000 组退化测试误差≤1e-12 m；错误的多烟幕入口会明确拒绝。",
        ),
        (
            "P4 结构证书",
            "保留 A 事件模型、B 解析上下界、C 向外舍入区间证书；G1 前提中显式写 8000 m 已锁定。",
            "A/B/C 可行性一致；正分离下界约 13.7916 s。",
        ),
        (
            "P5 参数化补偿族",
            "输出 t_m、云心 c*=s(t_m)、覆盖区间、起爆区间、实际释放与命令关系；不要只报一个代表坐标。",
            "起爆相对 t_m 的区间为 [h-18,-h]；宽度约 7.6239 s。",
        ),
        (
            "P6 场景输入与可执行筛选",
            "从 planning/scenario_schema.json 读取绝对初态和时钟；先检查命令合法，再检查释放点可达、航迹连续和 12 km 解释。",
            "输入不全时 T_executable_star=not_evaluated；输入完整时给出可复核约束残差。",
        ),
        (
            "P7 四类状态",
            "分别输出 execution_status、input_status、feasibility_status、certificate_status。",
            "不再使用一个 PASS 同时代表“程序成功”和“数学方案可行”。",
        ),
        (
            "P8 扩展边界",
            "G2 只检查必要条件；S2 用相对速度，包含衰减阶段时运行事件模型；无风数据只输出参数曲线。",
            "扩展结果不会覆盖 G1+S1+O0+U0 的名义结果文件。",
        ),
        (
            "P9 Q1→Q2 接口",
            "Q2 接收事件链和 Delta 定义，并检查舰船圆盘被烟幕并集覆盖；禁止使用 n×10.3761 作为联合上界。",
            "Q2 输出同时含 command 和 release 字段；旧别名带说明。",
        ),
    ]
    add_table(
        doc,
        ["任务", "具体怎么做", "验收标准"],
        programmer_tasks,
        [1500, 4860, 3000],
    )

    heading(doc, "5.3 编程手必须补的测试", 2)
    tests = [
        "□ 探测激活：G1 锁定后视轴偏差 β(t)=0；未启用“8000 m 已锁定”时不得沿用完整窗口结论。",
        "□ 时间链：t_d-t_cmd=2，t_b-t_d=3.5，t_b-t_cmd=5.5。",
        "□ 满半径区间：起爆位于 [t_m+h-18,t_m-h] 时整个最大覆盖段均处于 120 m 阶段。",
        "□ 缺陷退化：单烟幕时 single_smoke_margin=-Delta。",
        "□ 反例：把烟幕中心移远后验证器必须返回裸露，不能仍显示 PASS。",
        "□ 理想可达场景：输入完整且约束宽松时，T_executable_star 可达到结构上界。",
        "□ 不可达场景：无人机来不及到达释放点时，可执行值应严格下降或返回不可行。",
        "□ 12 km 边界：分别测试略小于、恰好等于、略大于 12000 m。",
        "□ G2 必要非充分：构造短窗口但无人机不可达的场景。",
        "□ S2 参数曲线：顺航、逆航、横向漂移只进入扩展输出。",
    ]
    for item in tests:
        bullet(doc, item)

    heading(doc, "六、论文手分工：你最终要写什么", 1)
    callout(
        doc,
        "论文手的核心职责",
        "把“先证明可行集为空，再求最优补偿”写成一条完整逻辑链。"
        "不要从旧乱码报告复制文字；所有数字从 JSON/CSV 取，所有结论带模型前提。",
        fill="EAF3E8",
    )

    heading(doc, "6.1 推荐的十段正文顺序", 2)
    writer_sections = [
        "问题重述：先判定全窗口持续完整遮蔽是否可行，不预设唯一最优解存在。",
        "条件分层：把题面常数、人工解释、简化假设、缺失场景输入分开。",
        "坐标系：地面坐标系用于绝对轨迹；舰船随动坐标系用于解释固定云心的相对运动。",
        "探测激活层：距离≤8000 m 且视轴偏差≤15°；G1 标准场景补充“8000 m 已锁定”。",
        "事件链：命令 t_cmd、实际释放 t_d、起爆 t_b；说明 2 s 端点是人工采用的主解释。",
        "完整覆盖层：定义 Delta(t)，再说明单烟幕时退化为旧裕度 g(t)=-Delta(t)。",
        "结构上界证明：推导 10.3761 s，并强调它不是烟幕 23 s 寿命。",
        "探测窗口与不可行证书：推导 24.1677 s 和 13.7916 s，说明 A/B/C 交叉验证。",
        "最大补偿族与可执行筛选：给 t_m、t_b、t_d、t_cmd、p_d 的参数关系，再区分结构上界与可执行最优。",
        "验证、局限和接口：写有界扰动、G2/S2 边界、缺失输入及 Q1→Q2 联合覆盖。",
    ]
    for item in writer_sections:
        numbered(doc, item, num_id)

    heading(doc, "6.2 论文手必须解释清楚的公式", 2)
    rows = [
        ("完整覆盖", "||s(t)-c||+R_s≤r(t)", "单烟幕必要且充分条件"),
        ("统一缺陷", "Delta=max_x min_j(||x-c_j||-r_j)", "Delta≤0 才完整覆盖"),
        ("结构上界", "2(R_c-R_s)/V_s", "固定单烟幕"),
        ("窗口下界", "(D_max-R_s)/(V_m+V_s)", "G1 且 8000 m 已锁定"),
        ("补偿中点", "t_m∈[t_in+h,t_out-h]", "不是命令时刻"),
        ("起爆区间", "t_b∈[t_m+h-18,t_m-h]", "保证满半径穿越"),
        ("事件链", "t_d=t_cmd+2; t_b=t_d+3.5", "主计时解释"),
        ("惯性位移", "c=p_d+98e_u", "模型假设，不是直接事实"),
    ]
    add_table(doc, ["用途", "公式", "写作提示"], rows, [1500, 4300, 3560])

    heading(doc, "6.3 论文手建议准备的图表", 2)
    rows = [
        ("图1", "地面坐标系与舰船随动坐标系", "固定烟幕相对舰船反向运动"),
        ("图2", "Delta(t) 或 g(t) 曲线", "标出遮蔽段、裸露段和临界点"),
        ("图3", "结构上界与探测窗口时间轴", "直观看出 10.3761<24.1677"),
        ("图4", "补偿族三种代表策略", "前段、中间均衡、后段保护"),
        ("表1", "题面条件/人工解释/缺失输入", "不要把假设写成事实"),
        ("表2", "A/B/C 证据对照", "正文 B 为主，A 复核，C 可放附录"),
        ("表3", "结构上界与可执行最优", "当前可执行值因缺输入未评估"),
        ("表4", "鲁棒性与翻转阈值", "参数范围必须注明探索性质"),
    ]
    add_table(doc, ["编号", "内容", "必须表达"], rows, [1000, 3400, 4960])

    heading(doc, "七、论文中允许和禁止的表述", 1)
    heading(doc, "7.1 可以写", 2)
    allowed = [
        "“在 G1+S1+O0+U0 且 8000 m 已锁定条件下，单弹全窗口持续完整遮蔽严格不可行。”",
        "“10.3761 s 是固定单烟幕的结构上界；可执行最优不超过该上界。”",
        "“由于绝对初态和任务时钟缺失，本问输出参数化最优补偿族，而非虚构唯一坐标。”",
        "“Q2 在 Q1 单烟幕接口上进一步考虑多烟幕的时序接力和空间互补。”",
    ]
    for item in allowed:
        bullet(doc, "✓ " + item)

    heading(doc, "7.2 不能写", 2)
    forbidden = [
        "“问题一无解。”——应说明全窗口方案不可行，但最大补偿族存在。",
        "“程序 PASS，所以防御方案可行。”——程序状态与数学可行性不同。",
        "“98 m 是题目直接给出的位移。”——它来自 S1 惯性假设。",
        "“固定航向窗口≤10.3761 s 就一定可行。”——这只是必要条件。",
        "“风漂使结果达到某数值。”——题面没有风数据。",
        "“单弹10.3761 s，所以三弹上界31.1284 s。”——联合覆盖可能有空间互补。",
        "“得到唯一最优坐标。”——绝对场景输入缺失且补偿族本身非唯一。",
    ]
    for item in forbidden:
        bullet(doc, "× " + item)

    heading(doc, "八、两名队员之间如何交接", 1)
    rows = [
        ("编程手→论文手", "q1_architecture_upgrade.json", "四类状态、核心数字、事件语义、缺陷测试"),
        ("编程手→论文手", "q1_parametric_compensation.json", "补偿族和达到上界的条件"),
        ("编程手→论文手", "q1_global_certificate.json", "向外舍入不可行证书"),
        ("编程手→论文手", "q1_robustness_summary.json", "扰动结果和扩展边界"),
        ("论文手→编程手", "公式与符号核对表", "确认 t_cmd、t_d、t_b、t_m 没有混用"),
        ("论文手→编程手", "待画图清单", "每张图列数据文件、横纵轴、要支持的结论"),
        ("双方→建模手", "阻塞项列表", "只请建模手判断真实需要判断的口径，不上交机械问题"),
    ]
    add_table(doc, ["方向", "交接物", "检查重点"], rows, [1900, 3000, 4460])

    heading(doc, "九、推荐工作节奏", 1)
    rhythm = [
        "第 1 轮（约 1 小时）：编程手复跑四轮证据；论文手同时阅读修订报告和接口合同。",
        "第 2 轮（约 2 小时）：编程手整理机器可读表；论文手写完前 8 段模型和证明。",
        "第 3 轮（约 2 小时）：编程手补图表数据、边界测试；论文手完成补偿族、验证和局限。",
        "第 4 轮（约 1 小时）：双方逐项核对数字、符号、单位、状态和文件路径。",
        "若获得真实初始坐标：另开场景化运行，不覆盖结构证书；生成 T_executable_star 和可达代表方案。",
    ]
    for item in rhythm:
        numbered(doc, item, num_id)

    heading(doc, "十、文件地图", 1)
    rows = [
        ("修订结论", "results/Q1/reports/q1_teammate_review_revision.md"),
        ("架构升级指标", "results/Q1/experiments/round4/metrics/q1_architecture_upgrade.json"),
        ("结构指标", "results/Q1/experiments/round1/metrics/q1_structural_metrics.json"),
        ("区间证书", "results/Q1/experiments/round2/metrics/q1_global_certificate.json"),
        ("补偿族", "results/Q1/experiments/round3/metrics/q1_parametric_compensation.json"),
        ("鲁棒性", "robustness/Q1/q1_robustness_summary.json"),
        ("场景输入规范", "planning/scenario_schema.json"),
        ("假设登记", "planning/assumption_register.csv"),
        ("Q1→Q2接口", "interfaces/Q1_to_Q2_coverage_contract.md"),
        ("代码审查", "code/Q1/reviews/q1_python_review_round4.json"),
        ("一致性审查", "results/Q1/reports/q1_q2_scoped_consistency_round4.json"),
    ]
    add_table(doc, ["用途", "相对路径"], rows, [2300, 7060])

    heading(doc, "十一、当前仍需建模手决定或补充的事项", 1)
    callout(
        doc,
        "真正的阻塞项",
        "若要输出唯一的可执行投放坐标，请提供舰船、导弹、无人机初始位置和航向、"
        "任务时钟零点，以及 12 km 到底相对起飞点、实时舰船还是总航程。"
        "在这些数据缺失时，结构性不可行证书和参数化补偿族仍然有效，但不能伪造绝对方案。",
        fill=CAUTION,
        color=RED,
    )

    doc.core_properties.title = "Q1 修订版编程手与论文手分工"
    doc.core_properties.subject = "依据队友总体审查结论升级的问题一执行手册"
    doc.core_properties.author = "建模组"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
