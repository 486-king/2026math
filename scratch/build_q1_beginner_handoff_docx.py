from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\qq\2026数模校赛")
OUT = ROOT / "Q1_编程手与中文手新手任务清单.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
PALE_GREEN = "EAF4EA"
PALE_GOLD = "FFF6DD"
PALE_RED = "FCEBEC"
WHITE = "FFFFFF"


def set_run_font(run, size=11, bold=None, italic=None, color="000000",
                 ascii_font="Calibri", east_font="Microsoft YaHei"):
    run.font.name = ascii_font
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    fonts.set(qn("w:eastAsia"), east_font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="B8C2CC", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_borders(table)


def style_cell_text(cell, bold=False, color="000000", size=9.3):
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.12
        for r in p.runs:
            set_run_font(r, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(text)
        set_cell_shading(cell, header_fill)
        style_cell_text(cell, bold=True, color=INK, size=font_size)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = str(value)
            style_cell_text(cell, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title, body, fill=CALLOUT, accent=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2, color="222222")
    set_table_geometry(table, [9360])
    set_table_borders(table, color=accent, size=8)
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def set_style_font(style, size, color, bold=False):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr()
    fonts = style._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, 11, "000000")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        set_style_font(style, size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        set_style_font(style, 11, "000000")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "CodeBlock" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("CodeBlock", 1)
    else:
        code_style = doc.styles["CodeBlock"]
    set_style_font(code_style, 9.3, INK)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.18)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(5)
    code_style.paragraph_format.line_spacing = 1.05


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_formula(doc, formula, note=None):
    p = doc.add_paragraph(style="CodeBlock")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(formula)
    set_run_font(r, size=10.2, color=INK, ascii_font="Cambria Math")
    if note:
        r2 = p.add_run(f"    {note}")
        set_run_font(r2, size=9.2, italic=True, color=MUTED)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    return p


def add_file(doc, path, purpose):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(path)
    set_run_font(r, size=9.4, bold=True, color=DARK_BLUE, ascii_font="Consolas")
    r2 = p.add_run(f" — {purpose}")
    set_run_font(r2, size=9.8, color="333333")


def add_check_table(doc, rows):
    return add_table(
        doc,
        ["状态", "要做的事", "完成标准"],
        [["待完成", a, b] for a, b in rows],
        [900, 5160, 3300],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("B题·问题一  |  团队执行手册")
    set_run_font(r, size=9, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("第 ")
    add_page_field(p)
    p.add_run(" 页")
    for run in p.runs:
        set_run_font(run, size=9, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(85)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("数学建模校赛 · B题 · 问题一")
    set_run_font(r, size=11, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("编程手与中文手\n新手任务清单")
    set_run_font(r, size=29, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("假设队友从未看过题目，也能按步骤完成交付")
    set_run_font(r, size=13.5, color=DARK_BLUE)

    add_callout(
        doc,
        "当前结论",
        "在建模手确认的 M1（纯追踪导弹）与 S1（起爆后烟幕中心固定）条件下，"
        "单枚干扰弹不可能在导弹的整个探测窗口内持续完整遮蔽舰船。"
        "正确答法不是继续硬搜一个不存在的坐标，而是先给出严格不可行证明，再给出最大连续遮蔽的参数化补偿方案。",
        fill=PALE_GREEN,
        accent=DARK_BLUE,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run(f"版本日期：{date.today().isoformat()}  |  使用状态：Q1 已完成并可交接")
    set_run_font(r, size=9.5, color=MUTED)
    doc.add_page_break()


def build():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_cover(doc)

    doc.add_heading("1. 先用白话看懂问题一", level=1)
    doc.add_paragraph(
        "一艘舰船正在匀速航行，一枚导弹接近舰船。无人机只能投放一枚干扰弹；"
        "干扰弹飞行 3.5 秒后起爆，形成烟幕。我们要判断：这一个烟幕能不能在导弹发现舰船后，"
        "一直到导弹接触舰船前，把整艘舰船完全挡住。如果做不到，就要严谨地说明为什么做不到，"
        "并给出“最多能连续挡多久、怎样布置才达到这个上限”。"
    )
    add_callout(
        doc,
        "最容易误解的地方",
        "烟幕从起爆到消失可持续 23 秒，并不代表它能完整遮住舰船 23 秒。"
        "烟幕中心固定，而舰船一直移动；只有舰船的整个半径 80 m 圆盘都落在烟幕圆盘内，才算“完整遮蔽”。",
        fill=PALE_GOLD,
        accent="7A5A00",
    )

    doc.add_heading("1.1 题目中的角色", level=2)
    add_table(
        doc,
        ["对象", "它怎么运动", "问题一里要关注什么"],
        [
            ["舰船", "速度 7.71 m/s，二维平面匀速直线航行", "舰船不是一个点，而是半径 80 m 的圆盘"],
            ["导弹", "M1：速度大小 320 m/s，方向实时指向舰船中心", "进入 8000 m 探测范围后，到接触舰船前的时间窗口"],
            ["无人机", "速度 28 m/s，作战半径 12000 m", "是否能在允许时间内到达投放点"],
            ["干扰弹", "投放后 3.5 s 内继承无人机投放瞬间的水平速度", "起爆点比投放点沿无人机方向前移 98 m"],
            ["烟幕", "起爆后中心固定；最大半径 120 m 保持 18 s，再用 5 s 衰减为 0", "能完整覆盖舰船圆盘的连续时段"],
        ],
        [1300, 3900, 4160],
    )

    doc.add_heading("1.2 已经由建模手确认的口径", level=2)
    add_bullet(doc, "M1（名义模型）：导弹做纯追踪，速度方向始终指向舰船中心。")
    add_bullet(doc, "M2（稳健性对照）：导弹保持固定航向；只有补齐初始位置与航向后才能计算。")
    add_bullet(doc, "S1（名义烟幕）：干扰弹惯性飞行 3.5 秒，起爆后烟幕中心固定。")
    add_bullet(doc, "名义模型不引入风漂移；漂移仅作为条件性扩展，不能写成题面事实。")
    add_bullet(doc, "目标是“全探测窗口内完整遮蔽舰船圆盘”，不是只遮住舰船中心，也不是最大化某几个采样时刻的覆盖率。")

    doc.add_heading("2. 当前已经证实的结论", level=1)
    add_formula(doc, "完整遮蔽条件：‖s(t) − c‖ + Rₛ ≤ r(t)")
    doc.add_paragraph(
        "烟幕最大半径为 120 m，舰船半径为 80 m，所以烟幕中心与舰船中心的距离最多只能是 40 m。"
        "舰船以 7.71 m/s 穿过半径 40 m 的允许圆，最长路径是直径 80 m。"
    )
    add_formula(doc, "T_cover^max = 2(R_c − R_s)/V_s = 80/7.71 = 10.3761348898 s")
    doc.add_paragraph(
        "M1 纯追踪下，导弹从距离 8000 m 进入探测，到接触舰船半径 80 m，"
        "即使按可能的最大闭合速度 320+7.71 m/s 计算，也至少需要："
    )
    add_formula(doc, "T_detect^min = (8000 − 80)/(320 + 7.71) = 24.1677092551 s")
    add_formula(doc, "T_naked^min = T_detect^min − T_cover^max = 13.7915743654 s > 0")
    add_callout(
        doc,
        "正式结论",
        "在一枚干扰弹、M1 纯追踪、S1 固定烟幕中心和题面常数下，全程完整遮蔽严格不可行。"
        "至少存在 13.7916 s 的总裸露时间下界。A 事件模型、B 解析 baseline 与 C 向外舍入证书结论一致。",
        fill=PALE_GREEN,
        accent="1F3A5F",
    )

    doc.add_heading("2.1 这条结论意味着什么", level=2)
    add_bullet(doc, "不是程序没找到解，而是所有可能的投放位置和时间都无法突破物理时长矛盾。")
    add_bullet(doc, "不能编造一个“最优绝对坐标”；题目缺少绝对初始位置，而且严格全程遮蔽可行集为空。")
    add_bullet(doc, "仍然可以求最优补偿：让烟幕达到 10.3761 s 的最大连续完整遮蔽，并按次级目标选择放在窗口前段、中间或后段。")
    add_bullet(doc, "如果以后采用 M2 或加入实际风漂移，必须重新计算，不能直接沿用名义结论。")

    doc.add_heading("3. 两位队友怎样配合", level=1)
    add_table(
        doc,
        ["角色", "主要责任", "不能擅自做的事", "交给下一位什么"],
        [
            ["编程手", "复现 A/B/C 证据；补齐测试；把结果保存为可追溯文件；有新场景数据时再做坐标计算", "不能用缺省 0 代替缺失坐标；不能只凭网格搜索声称连续时间全局结论", "机器可读结果、运行记录、验证表、作图数据"],
            ["中文手（论文写作）", "把已验证的公式、证明、补偿族和适用范围写清楚", "不能发明坐标、风速、实验结果或引用；不能把条件结论写成普遍结论", "问题一完整文字稿、公式编号建议、图表需求单"],
            ["建模手", "确认数学口径、符号、结论范围与最终选用的代表补偿方案", "不能让论文数字与代码结果脱节", "最终审批意见与需要继续扩展的范围"],
        ],
        [1300, 3300, 2900, 1860],
        font_size=8.5,
    )

    doc.add_heading("4. 编程手：从零开始的工作清单", level=1)
    add_callout(
        doc,
        "你的最终任务",
        "让任何队友在同一工作区运行程序后，都能得到同样的三个时间量和同样的不可行结论；"
        "同时为论文手提供可直接画图、制表和核对的结果文件。现在不需要重写主模型。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )

    doc.add_heading("4.1 第一步：只读这 6 类文件", level=2)
    add_file(doc, "B题：舰船烟幕遮蔽干扰优化.docx", "题目原文；只读，不覆盖")
    add_file(doc, "methods/Q1/q1_final_method_explanation.md", "最终采用的数学模型与求解逻辑")
    add_file(doc, "results/Q1/reports/q1_final_result_analysis.md", "正式数值、补偿方案和限制")
    add_file(doc, "code/Q1/q1_code_plan.md", "代码输入、输出、模块和验收合同")
    add_file(doc, "planning/symbol_table.md", "全队统一符号与单位")
    add_file(doc, "methods/Q1/q1_decisions.jsonl", "建模手亲自确认的选择和口径")

    doc.add_heading("4.2 第二步：先复现已经完成的结果", level=2)
    doc.add_paragraph("在项目根目录打开终端，依次运行：")
    for cmd in (
        "python code/Q1/q1_run.py",
        "python code/Q1/q1_interval_certificate.py",
        "python code/Q1/q1_robustness.py",
    ):
        p = doc.add_paragraph(style="CodeBlock")
        p.add_run(cmd)
    add_check_table(
        doc,
        [
            ("三个命令都能正常结束，不出现异常堆栈。", "退出状态正常；对应结果文件时间戳更新或被一致地复核。"),
            ("A 与 B 返回相同的三项核心结果。", "差值不超过 10⁻⁹ s；comparison.status 为 PASS。"),
            ("C 返回严格的正分离区间。", "certificate_status 为 PASS，positive_separation_lower_s > 0。"),
            ("程序没有生成虚构坐标。", "coordinate_solution_status 明确说明无全程可行坐标，或缺少场景输入。"),
        ],
    )

    doc.add_heading("4.3 第三步：知道每个程序负责什么", level=2)
    add_table(
        doc,
        ["文件", "小白版说明", "你主要检查什么"],
        [
            ["q1_common.py", "统一保存常数、烟幕半径函数和通用公式", "单位、边界值、所有模块是否调用同一组常数"],
            ["q1_event_model.py", "方案 A：按探测、起爆、半径变化等事件检查连续时间", "不能只检查有限网格点；事件边界和覆盖根不能漏"],
            ["q1_analytic_baseline.py", "方案 B：用几何和速度界直接算理论上界/下界", "公式必须与论文完全一致"],
            ["q1_interval_certificate.py", "方案 C：用向外舍入确保浮点误差也推翻不了正间隔", "证书状态与区间方向正确"],
            ["q1_run.py", "统一入口，运行 A/B 并保存比较结果", "结果文件路径、状态、警告和环境信息完整"],
            ["q1_robustness.py", "做参数扰动、M2 条件阈值和漂移边界检查", "扩展结果不能混进名义结论"],
        ],
        [2100, 4300, 2960],
        font_size=8.6,
    )

    doc.add_heading("4.4 第四步：补齐必须通过的测试", level=2)
    add_check_table(
        doc,
        [
            ("烟幕半径分段函数边界", "r(0)=120、r(18)=120、r(23)=0；区间外为 0。"),
            ("投放—起爆关系", "惯性位移为 28×3.5=98 m；p_d=c−98e_u。"),
            ("共线追踪人造特例", "数值积分恢复解析接触时间，误差不超过既定容差。"),
            ("烟幕中心在舰船航迹上", "恢复 80/7.71=10.3761348898 s。"),
            ("烟幕中心偏离舰船航迹", "连续完整遮蔽时长严格小于理论上界。"),
            ("A/B/C 一致性", "三种方法在名义口径下均返回 strict_full_window_feasible=false。"),
            ("单位检查", "位置 m、速度 m/s、时间 s、角度 rad；15°只在输入或展示处保留。"),
            ("输入缺失检查", "缺少初始坐标时返回结构化阻塞状态，绝不自动补零。"),
        ],
    )

    doc.add_heading("4.5 如果后来拿到真实初始坐标，照这个顺序扩展", level=2)
    for text in (
        "建立场景输入文件，至少含舰船初始位置与航向、导弹初始位置、无人机初始位置、任务时钟定义；M2 还要固定导弹航向。",
        "先运行 M1 轨迹，求出实际探测进入时刻 t_in 和接触/退出时刻 t_out。",
        "选择最大连续遮蔽的中心时刻 t_c，并令烟幕中心 c*=s(t_c)。",
        "在允许起爆区间内寻找 t_b，再由 t_d=t_b−3.5 和 p_d=c*−98e_u 得到投放事件。",
        "逐项检查响应延时 t_d≥2 s、无人机 28 m/s 可达、12000 m 作战半径和真实航迹可执行性。",
        "若理论上界不可达，报告实际最优值及它与 10.3761 s 理论上界的差距；不要隐去约束失败。",
        "保存轨迹、探测窗口、覆盖裕度和裸露区间数据，不只保存一个最优值。",
    ):
        add_number(doc, text)

    doc.add_heading("4.6 M2 与风漂移什么时候才做", level=2)
    add_table(
        doc,
        ["扩展", "启动条件", "需要的新输入", "判断标准"],
        [
            ["M2 固定航向", "拿到导弹初始位置和固定航向；或建模手要求做数值对照", "导弹初始坐标、航向、统一时钟", "实际距离—视场探测窗口是否 ≤10.3761 s；若结论改变，必须重新报建模手"],
            ["烟幕漂移", "题目补充风向风速，或建模手批准鲁棒性范围", "漂移速度向量及其依据", "用舰船—烟幕相对速度重算覆盖上界；不得把 +5 m/s 探索值当实测风"],
        ],
        [1550, 2600, 2400, 2810],
        font_size=8.5,
    )

    doc.add_heading("4.7 编程手最后要交什么", level=2)
    add_check_table(
        doc,
        [
            ("可一键复现的程序与说明", "新队友只看 README/命令即可得到结果。"),
            ("A/B/C 运行摘要", "每种方法的状态、核心数值、时间、警告和输入来源齐全。"),
            ("测试报告", "上述边界、特例、一致性和缺输入测试全部有记录。"),
            ("论文作图数据", "CSV/JSON 明确字段、单位和来源；不从截图抄数。"),
            ("变更说明", "如果改公式、常数或口径，先通知建模手并同步所有证据文件。"),
        ],
    )
    add_callout(
        doc,
        "编程手禁区",
        "不要随机造初始坐标；不要把缺失值置零；不要只用离散网格证明连续可行性；"
        "不要把 M2 或风漂移的条件性结果覆盖名义 M1/S1 结果；不要为了得到“漂亮最优解”删掉硬约束。",
        fill=PALE_RED,
        accent="9B1C1C",
    )

    doc.add_heading("5. 中文手（论文写作）：从零开始的工作清单", level=1)
    add_callout(
        doc,
        "你的最终任务",
        "把“为什么一枚烟幕做不到全程遮蔽”和“做不到以后怎样达到最大连续遮蔽”写成读者能复核的完整论证。"
        "所有数字必须来自已保存结果，所有条件必须写清。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )

    doc.add_heading("5.1 动笔前先读这些文件", level=2)
    add_file(doc, "B题：舰船烟幕遮蔽干扰优化.docx", "先理解题目问的对象与限制")
    add_file(doc, "methods/Q1/q1_final_method_explanation.md", "论文方法部分的唯一主口径")
    add_file(doc, "results/Q1/reports/q1_final_result_analysis.md", "结果、补偿族、达到条件和适用范围")
    add_file(doc, "planning/symbol_table.md", "符号、含义和单位")
    add_file(doc, "planning/model_assumptions.md", "哪些是题面条件，哪些是模型假设")
    add_file(doc, "robustness/Q1/q1_robustness_summary.json", "稳健性边界；只取已验证内容")

    doc.add_heading("5.2 推荐的论文叙事顺序", level=2)
    add_table(
        doc,
        ["顺序", "这一小节写什么", "读者看完应明白什么"],
        [
            ["1", "问题重述与目标", "我们先检验单弹全程完整遮蔽是否可行，而不是默认存在最优坐标"],
            ["2", "坐标、对象、符号与假设", "舰船、导弹、无人机、干扰弹和烟幕各怎样运动；M1/S1 的适用范围"],
            ["3", "完整遮蔽判据", "为什么必须满足 ‖s−c‖+R_s≤r，而不是只比较两个中心"],
            ["4", "单烟幕覆盖上界", "固定烟幕最多完整遮蔽 10.3761 s"],
            ["5", "M1 探测窗口下界", "纯追踪窗口至少 24.1677 s"],
            ["6", "全局不可行证明", "两个严格时间界相差至少 13.7916 s，所以可行集为空"],
            ["7", "不可行后的最优补偿", "怎样放置一个长度为 10.3761 s 的最大连续遮蔽区间"],
            ["8", "A/B/C 验证与稳健性", "结论不是局部优化、网格或浮点误差造成；哪些改动可能改变结论"],
            ["9", "问题一结论与对问题二的接口", "单弹只能提供有限遮蔽区间，多弹模型应做区间衔接"],
        ],
        [700, 3770, 4890],
        font_size=8.6,
    )

    doc.add_heading("5.3 方法部分必须出现的公式", level=2)
    add_formula(doc, "s(t)=s₀+Vₛeₛt")
    add_formula(doc, "ṁ(t)=Vₘ [s(t)−m(t)]/‖s(t)−m(t)‖", "M1 纯追踪")
    add_formula(doc, "t_b=t_d+3.5，c=p_d+3.5V_ue_u=p_d+98e_u")
    add_formula(doc, "g(t)=r(t)−Rₛ−‖c−s(t)‖；完整遮蔽要求 g(t)≥0")
    add_formula(doc, "T_cover^max=2(R_c−R_s)/V_s=10.3761348898 s")
    add_formula(doc, "T_detect≥(D_max−R_s)/(V_m+V_s)=24.1677092551 s")
    add_formula(doc, "T_naked≥T_detect−T_cover^max=13.7915743654 s")
    doc.add_paragraph(
        "写公式时要在正文中解释每个量的单位和物理意义。不要连续堆公式后只写“显然可得”。"
    )

    doc.add_heading("5.4 不可行以后，补偿方案要这样写", level=2)
    doc.add_paragraph(
        "令 h=(R_c−R_s)/V_s=5.1880674449 s，实际探测窗口为 [t_in,t_out]。"
        "为了达到 2h=10.3761348898 s 的最大连续完整遮蔽，烟幕中心必须落在舰船航迹上。"
    )
    add_formula(doc, "t_c∈[t_in+h, t_out−h]，c*=s(t_c)")
    add_formula(doc, "I_cover=[t_c−h, t_c+h]")
    add_formula(doc, "t_b∈[t_c+h−18, t_c−h]")
    add_formula(doc, "t_d=t_b−3.5，p_d=c*−98e_u，‖e_u‖=1")
    add_bullet(doc, "若想让两端不可避免的裸露尽量均衡，取 t_c=(t_in+t_out)/2。")
    add_bullet(doc, "若优先从探测开始就保护，取 t_c=t_in+h。")
    add_bullet(doc, "若强调尽量晚起爆，取 t_b=t_c−h。")
    add_bullet(doc, "任何代表解都必须再检查无人机响应延时、可达性和 12 km 作战半径。")

    doc.add_heading("5.5 哪些句子可以写，哪些不能写", level=2)
    add_table(
        doc,
        ["可以这样写", "不要这样写", "原因"],
        [
            ["在 M1/S1 和题面常数下，单弹全程完整遮蔽严格不可行。", "问题一无解。", "前者写清适用条件；后者把条件结论扩大成普遍结论。"],
            ["全局可行集为空，原因是探测窗口下界大于覆盖上界。", "优化器没有搜到解。", "不可行来自证明，不是求解器表现。"],
            ["最大连续完整遮蔽上界为 10.3761 s。", "烟幕只能存在 10.3761 s。", "烟幕可存在 23 s，但完整遮蔽时长更短。"],
            ["至少存在 13.7916 s 的总裸露时间下界。", "裸露时间恰好是 13.7916 s。", "当前是由最短探测窗口得到的下界，不是所有场景的精确值。"],
            ["绝对投放坐标需补充初始几何后计算。", "最优投放坐标为（0,0）。", "题面没有支持该坐标，不能编造。"],
            ["M2 或足够强的顺航向漂移可能改变时长必要条件。", "考虑风后仍然一定不可行。", "没有风数据，扩展只能条件性表述。"],
        ],
        [3200, 2800, 3360],
        font_size=8.3,
    )

    doc.add_heading("5.6 建议向编程手索要的图表", level=2)
    add_check_table(
        doc,
        [
            ("时间轴图", "同一横轴画出 M1 探测窗口下界和 10.3761 s 最大遮蔽区间，直观看见长度矛盾。"),
            ("几何示意图", "舰船圆盘半径 80 m、烟幕圆盘半径 120 m、允许中心偏差 40 m、舰船直线穿越。"),
            ("A/B/C 对照表", "列出覆盖上界、探测下界、裸露下界、可行性判断和验证方式。"),
            ("适用范围表", "名义 M1/S1、M2 条件、烟幕漂移扩展分别需要什么输入、结论能否沿用。"),
        ],
    )
    doc.add_paragraph(
        "图表必须从 CSV/JSON 或已验证公式生成。探索图、终端截图和未经复核的临时表不能直接放进论文。"
    )

    doc.add_heading("5.7 中文手交稿前逐项核对", level=2)
    add_check_table(
        doc,
        [
            ("题意完整", "回答了“能否全程遮蔽”；不可行后又给出最大连续遮蔽和投放参数关系。"),
            ("条件完整", "一枚干扰弹、M1、S1、无名义风漂移、完整圆盘遮蔽全部写明。"),
            ("数字可追溯", "10.3761、24.1677、13.7916 和 98 m 均与结果文件一致。"),
            ("上下界用词准确", "下界不写成精确值，上界不写成实际必达值；达到条件单列。"),
            ("符号一致", "全文与 planning/symbol_table.md 一致，没有同符号多义。"),
            ("没有编造", "没有虚构坐标、风速、M2 数值结果、引用或奖项承诺。"),
            ("验证写清", "说明 A/B 一致、C 向外舍入证书、单因素 ±10% 结果及其限制。"),
            ("逻辑闭环", "从完整遮蔽判据到时间矛盾，再到补偿族，中间没有跳步。"),
        ],
    )

    doc.add_heading("6. 两位队友共同使用的验收表", level=1)
    add_table(
        doc,
        ["检查项", "编程手回答", "中文手回答", "最终通过条件"],
        [
            ["核心结论", "程序是否返回 M1/S1 严格不可行？", "正文是否明确写成条件性全局不可行？", "两边表述一致"],
            ["三个时间量", "JSON/CSV 数值与容差是否正确？", "公式、单位、小数位和上下界用词是否正确？", "能逐项追溯"],
            ["补偿方案", "参数化关系和可达性接口是否保留？", "是否说明多解与代表解选择？", "不虚构唯一坐标"],
            ["验证", "A/B/C、特例和边界测试是否通过？", "是否解释三种验证各排除了什么风险？", "证据与文字对应"],
            ["扩展", "M2/漂移是否与名义结果分开保存？", "是否只作条件性讨论？", "不混淆模型范围"],
            ["交接", "能否一键运行并给出结构化状态？", "能否让没看题的人复述证明链？", "另一位队友无需猜测"],
        ],
        [1600, 2600, 2600, 2560],
        font_size=8.4,
    )

    doc.add_heading("7. 常见问题与出错处理", level=1)
    add_table(
        doc,
        ["遇到的情况", "先做什么", "不要做什么"],
        [
            ["运行结果与 10.3761/24.1677/13.7916 不一致", "检查常数、单位、舰船半径是否从探测距离中扣除；对照 B 解析公式", "不要先调优化器参数掩盖差异"],
            ["程序要求初始坐标但题目没有", "返回 blocked_missing_scenario_inputs，并只运行不依赖坐标的结构证书", "不要补成零或随机数"],
            ["M2 得到可能可行", "保存实际探测窗口与输入，通知建模手重新判定结论范围", "不要直接替换正式 M1 结果"],
            ["加入漂移后结论翻转", "核对漂移方向、单位和数据来源，作为新模型重跑 A/B", "不要把探索用 +5 m/s 当题面数据"],
            ["论文想要一个具体坐标", "提供参数化关系和缺失输入清单；有真实坐标后再算", "不要为了版面完整编造坐标"],
            ["网格上所有点都可行/不可行", "回到连续事件、解析界或区间证书复核", "不要用有限网格冒充连续证明"],
        ],
        [2250, 4300, 2810],
        font_size=8.4,
    )

    doc.add_heading("8. 交付文件地图", level=1)
    add_table(
        doc,
        ["用途", "权威文件"],
        [
            ["题目原文", "B题：舰船烟幕遮蔽干扰优化.docx"],
            ["最终数学方法", "methods/Q1/q1_final_method_explanation.md"],
            ["最终结果解释", "results/Q1/reports/q1_final_result_analysis.md"],
            ["建模手选择记录", "methods/Q1/q1_decisions.jsonl"],
            ["代码实现合同", "code/Q1/q1_code_plan.md"],
            ["代码审查", "code/Q1/reviews/q1_python_review.json"],
            ["A/B 运行摘要", "results/Q1/experiments/round1/run_summary.json"],
            ["C 全局证书", "results/Q1/experiments/round2/metrics/q1_global_certificate.json"],
            ["补偿族结果", "results/Q1/experiments/round3/metrics/q1_parametric_compensation.json"],
            ["稳健性证据", "robustness/Q1/q1_robustness_summary.json"],
            ["原简版编程清单", "handoff/Q1/q1_programmer_checklist.md"],
            ["原简版写作清单", "handoff/Q1/q1_writer_checklist.md"],
        ],
        [2200, 7160],
        font_size=8.7,
    )

    doc.add_heading("9. 一句话交接模板", level=1)
    add_callout(
        doc,
        "编程手完成后这样汇报",
        "“我已复现 Q1 的 A/B/C 结果：最大连续完整遮蔽 10.3761348898 s，"
        "M1 探测窗口下界 24.1677092551 s，裸露下界 13.7915743654 s；"
        "一致性和边界测试已通过。当前没有使用虚构初始坐标。新增或未通过项目如下：……”",
        fill=LIGHT_GRAY,
        accent=DARK_BLUE,
    )
    add_callout(
        doc,
        "中文手完成后这样汇报",
        "“我已按‘可行性检查—全局不可行证明—最优补偿—验证与边界’写完 Q1；"
        "所有数字和符号均已对照权威文件，没有把下界写成精确值，也没有虚构坐标、风数据或 M2 结果。"
        "仍需建模手确认的表述如下：……”",
        fill=LIGHT_GRAY,
        accent=DARK_BLUE,
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— 问题一交接手册结束 —")
    set_run_font(r, size=10, bold=True, color=MUTED)

    # Core properties
    doc.core_properties.title = "B题问题一：编程手与中文手新手任务清单"
    doc.core_properties.subject = "数学建模问题一团队交接与验收手册"
    doc.core_properties.author = "数学建模团队"
    doc.core_properties.keywords = "数学建模, 问题一, 编程手, 中文手, 任务清单, M1, S1"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
